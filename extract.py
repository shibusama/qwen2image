"""从文档文件或 URL 提取纯文本。"""

import re

MAX_CHARS = 8000


def _truncate(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n[内容过长，已截断]"
    return text


def extract_pdf(data: bytes) -> str:
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(data)
    parts = []
    for page in pdf:
        parts.append(page.get_textpage().get_text_range())
    pdf.close()
    return _truncate("\n".join(parts))


def extract_docx(data: bytes) -> str:
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return _truncate("\n".join(parts))


def extract_plain(data: bytes) -> str:
    try:
        return _truncate(data.decode("utf-8"))
    except UnicodeDecodeError:
        return _truncate(data.decode("gbk", errors="ignore"))


def extract_url(url: str) -> str:
    import requests

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/120 Safari/537.36"
    }
    resp = requests.get(url, headers=headers, timeout=20)
    resp.raise_for_status()
    html = resp.text
    html = re.sub(r"(?is)<(script|style|noscript).*?</\1>", " ", html)
    html = re.sub(r"(?is)<br\s*/?>", "\n", html)
    html = re.sub(r"(?is)<[^>]+>", " ", html)
    html = re.sub(r"&nbsp;?", " ", html)
    html = re.sub(r"&amp;?", "&", html)
    html = re.sub(r"&lt;?", "<", html)
    html = re.sub(r"&gt;?", ">", html)
    return _truncate(html)


def extract_file_bytes(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_pdf(data)
    if name.endswith(".docx"):
        return extract_docx(data)
    if name.endswith(".txt") or name.endswith(".md"):
        return extract_plain(data)
    raise ValueError(f"不支持的文件类型：{filename}（支持 PDF / DOCX / TXT / MD）")


def image_size(raw: bytes):
    """从图片文件头解析原始宽高（PNG/JPEG，纯标准库）。解析失败返回 None。"""
    import struct

    try:
        if raw[:8] == b"\x89PNG\r\n\x1a\n":
            w, h = struct.unpack(">II", raw[16:24])
            return w, h
        if raw[:2] == b"\xff\xd8":  # JPEG
            i = 2
            while i + 9 < len(raw):
                if raw[i] != 0xFF:
                    i += 1
                    continue
                marker = raw[i + 1]
                if marker in (0xC0, 0xC1, 0xC2, 0xC3, 0xC5, 0xC6, 0xC7, 0xC9, 0xCA, 0xCB, 0xCD, 0xCE, 0xCF):
                    h, w = struct.unpack(">HH", raw[i + 5 : i + 9])
                    return w, h
                seg_len = struct.unpack(">H", raw[i + 2 : i + 4])[0]
                i += 2 + seg_len
    except Exception:
        pass
    return None
